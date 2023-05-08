import re
import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import tempfile
from pathlib import Path
import shutil


def extract_info_from_pdf(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        first_page = pdf.pages[0]
        text = first_page.extract_text()
        print(f"text: {text}")  # Print the extracted text

        landlord_name = re.search(r'\(1\)(.*),', text).group(1).strip()
        tenant_name = re.search(r'\(2\)(.*),', text).group(1).strip()
        address = re.search(r'adress(.*),', text).group(1).strip()
        transaction_search = re.search(r'Transaktion\s+(\S+)', text)
        transaction_id = transaction_search.group(1).strip() if transaction_search is not None else None


        current_rent = None
        for page in pdf.pages:
            text = page.extract_text()
            match = re.search(r'Hyran är\s+([\d\s]+)', text)
            if match:
                current_rent = match.group(1).strip()
                break

        if current_rent is None:
            raise ValueError("Current rent not found in the PDF")

    return landlord_name, tenant_name, address, transaction_id, current_rent

def replace_placeholders(doc, placeholder, value):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, value).replace('[', '').replace(']', '')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, value).replace('[', '').replace(']', '')

def create_new_rent_increase_pdf(template_path, landlord_name, tenant_name, application_date, current_rent, new_rent, service_fee, address, transaction_id, output_path):
    # Open the template file
    doc = Document(template_path)

    # Replace the placeholders with the extracted data and input variables
    replace_placeholders(doc, 'LANDLORD_NAME', landlord_name)
    replace_placeholders(doc, 'TENANT_NAME', tenant_name)
    replace_placeholders(doc, 'ADDRESS', address)
    replace_placeholders(doc, 'TRANSACTION_ID', transaction_id)
    replace_placeholders(doc, 'CURRENT_RENT', current_rent)
    replace_placeholders(doc, 'NEW_RENT', str(new_rent))  # Convert new_rent to string
    replace_placeholders(doc, 'SERVICE_FEE', str(service_fee))  # Convert service_fee to string
    replace_placeholders(doc, 'APPLICATION_DATE', application_date.strftime('%Y-%m-%d'))
    replace_placeholders(doc, 'TODAYS_DATE', datetime.today().strftime('%Y-%m-%d'))

    # Create the directory if it doesn't exist
    output_dir = os.path.dirname(output_path)
    os.makedirs(output_dir, exist_ok=True)

    # Save the document to the output_path
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = os.path.join(tmpdir, 'Rent_Increase.docx')
        doc.save(tmp_path)
        with open(tmp_path, 'rb') as f:
            with open(output_path, 'wb') as output_file:
                output_file.write(f.read())



downloads_folder = os.path.expanduser('~/Downloads')

# ... (previous code remains the same)

# Extract the required information from the lease PDF
"""pdf_path = 'uploaded_files/Hyresavtal för Täckaregatan, Åstorp.pdf'
landlord_name, tenant_name, address, transaction_id, current_rent = extract_info_from_pdf(pdf_path)

# Define the output file name and path
output_file_name = f'Rent_Increase_{transaction_id}.docx'
output_path = os.path.join(downloads_folder, output_file_name)

# Create a new rent increase PDF using the provided template
template_path = 'template.docx'
new_rent = int(input("Enter the new rent: "))
application_date = datetime.strptime(input("Enter the application date (yyyy-mm-dd): "), '%Y-%m-%d')

# Calculate service fee
service_fee = new_rent * 0.0495

create_new_rent_increase_pdf(template_path, landlord_name, tenant_name, application_date, new_rent, service_fee, address, transaction_id, output_path)"""
#extract the required information from the Lease PDF